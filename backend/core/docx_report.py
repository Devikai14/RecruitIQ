"""
HR Candidate Profiles document generator — with Resume Quality Score section.
"""
import io
import os
import json
import urllib.request
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_para(doc, text, size=10, bold=False, italic=False,
              color=(0x22, 0x22, 0x22), align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


def _quality_color(grade):
    return {
        "Excellent": "1B8A4A",
        "Good":      "E6900A",
        "Fair":      "CC5200",
        "Poor":      "CC0000",
    }.get(grade, "777777")


# ── LLM Summary ───────────────────────────────────────────────────────────────

def _generate_llm_summary(candidate):
    api_key  = os.environ.get("GROQ_API_KEY", "")
    name     = candidate.get("name", "Candidate")
    skills   = ", ".join(candidate.get("matched_skills", [])) or "Not specified"
    exp      = candidate.get("experience", 0)
    score    = candidate.get("score", "N/A")
    projects = candidate.get("projects", [])
    proj_str = " | ".join(projects[:2]) if projects and projects[0] != "No specific projects extracted" else "Not specified"

    prompt = (
        "You are an expert HR analyst. Write a concise, professional 3-4 sentence candidate summary "
        "for an HR interviewer. Be specific, write in third person, and make it interview-ready.\n\n"
        "Candidate: " + name + "\n"
        "Match Score: " + str(score) + "\n"
        "Years of Experience: " + str(exp) + "\n"
        "Key Skills: " + skills + "\n"
        "Notable Projects: " + proj_str + "\n\n"
        "Write ONLY the summary paragraph. Do not mention any score as a fraction or 'out of 100'."
    )

    if not api_key:
        return _rule_based_summary(candidate)

    try:
        payload = json.dumps({
            "model": "llama3-8b-8192",
            "max_tokens": 300,
            "messages": [{"role": "user", "content": prompt}]
        }).encode("utf-8")
        req = urllib.request.Request(
            "https://api.groqcom/openai/v1/chat/completions",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": "Bearer " + api_key,
            },
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print("Groq API error:", e)
        return _rule_based_summary(candidate)


def _rule_based_summary(candidate):
    name      = candidate.get("name", "Candidate")
    skills    = candidate.get("matched_skills", [])
    exp       = candidate.get("experience", 0)
    score     = candidate.get("score", "N/A")
    proj      = candidate.get("projects", [])
    skill_str = ", ".join(skills[:4]) if skills else "various technologies"
    proj_hint = ""
    if proj and proj[0] != "No specific projects extracted":
        proj_hint = " Notable work includes: " + proj[0][:80] + "."
    return (
        name + " is a professional with " + str(exp) + " year(s) of experience "
        "specialising in " + skill_str + ". "
        "The candidate achieved a match score of " + str(score) + " against the job requirements." +
        proj_hint + " This candidate is recommended for a structured technical interview."
    )


# ── Main document generator ───────────────────────────────────────────────────

def generate_hr_docx(top_candidates, interview_date, hr_name):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("RecruitIQ — Candidate Profiles Report")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    try:
        fmt_date = datetime.strptime(interview_date, "%Y-%m-%d").strftime("%A, %d %B %Y")
    except Exception:
        fmt_date = interview_date

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    sr = sub.add_run(f"Interview Date: {fmt_date}   |   Prepared for: {hr_name}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(14)
    nr = note.add_run("AI-assisted screening report — human review recommended before final decisions")
    nr.italic = True
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_after = Pt(14)
    hr_run = hr_para.add_run("─" * 95)
    hr_run.font.size = Pt(7)
    hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── Section 1: Schedule Table ─────────────────────────────────────────────
    sec1 = doc.add_paragraph()
    sec1.paragraph_format.space_after = Pt(8)
    s1r = sec1.add_run("SECTION 1 — Interview Schedule")
    s1r.bold = True
    s1r.font.size = Pt(11)
    s1r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    cols   = ["#", "Candidate Name", "Match Score", "Skills", "Experience", "Quality", "Slot"]
    widths = [0.5, 3.8, 1.6, 1.8, 1.6, 1.8, 3.4]
    table  = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, w in enumerate(widths):
        table.columns[i].width = int(Cm(w))

    hrow = table.rows[0]
    for i, col in enumerate(cols):
        cell = hrow.cells[i]
        cell.text = col
        _set_cell_bg(cell, "D6E4F0")
        _set_cell_border(cell, "AAAAAA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    for i, c in enumerate(top_candidates):
        qs    = c.get("quality_score", "—")
        qg    = c.get("quality_grade", "")
        skill_str = str(len(c.get("matched_skills", []))) + "/" + str(c.get("total_skills", "?"))
        vals = [
            str(i + 1),
            c.get("name", "?"),
            str(c.get("score", "?")),
            skill_str,
            str(c.get("experience", 0)) + " yr",
            f"{qs}/100 ({qg})" if qs != "—" else "—",
            c.get("slot", "TBD"),
        ]
        bg  = "FFFFFF" if i % 2 == 0 else "F5F7FA"
        row = table.add_row()
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            _set_cell_bg(cell, bg)
            _set_cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(9)
            # Colour the quality cell by grade
            if j == 5 and qg:
                hx = _quality_color(qg)
                run.font.color.rgb = RGBColor(
                    int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
                )
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    doc.add_paragraph()

    # ── Section 2: Detailed Profiles ──────────────────────────────────────────
    sec2 = doc.add_paragraph()
    sec2.paragraph_format.space_after = Pt(10)
    s2r = sec2.add_run("SECTION 2 — Detailed Candidate Profiles")
    s2r.bold = True
    s2r.font.size = Pt(11)
    s2r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    for i, c in enumerate(top_candidates):
        rank_p = doc.add_paragraph()
        rank_p.paragraph_format.space_after = Pt(4)
        rr = rank_p.add_run(f"#{i + 1}  {c.get('name', 'Unknown')}")
        rr.bold = True
        rr.font.size = Pt(13)
        rr.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

        # AI Summary
        summ_lbl = doc.add_paragraph()
        summ_lbl.paragraph_format.space_after = Pt(3)
        sl = summ_lbl.add_run("AI Summary")
        sl.bold = True
        sl.font.size = Pt(9)
        sl.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        ai_summary = _generate_llm_summary(c)
        summ_p = doc.add_paragraph()
        summ_p.paragraph_format.space_after = Pt(8)
        summ_r = summ_p.add_run(ai_summary)
        summ_r.italic = True
        summ_r.font.size = Pt(9.5)
        summ_r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Profile info table
        matched_skills = c.get("matched_skills", [])
        all_skills     = c.get("skill_list", [])
        missing_skills = [s for s in all_skills if s not in matched_skills]
        projects       = c.get("projects", ["N/A"])
        proj_str       = " | ".join(projects[:3]) if projects else "N/A"

        qs    = c.get("quality_score", "N/A")
        qg    = c.get("quality_grade", "")
        qpass = ", ".join(c.get("quality_passed", [])) or "None"
        qfail = ", ".join(c.get("quality_failed", [])) or "None"

        fields = [
            ("Match Score",    str(c.get("score", "?"))),
            ("Experience",     str(c.get("experience", 0)) + " year(s)"),
            ("Email",          c.get("email") or "Not found in resume"),
            ("Interview Slot", c.get("slot", "TBD")),
            ("Skills Matched", ", ".join(matched_skills) if matched_skills else "None"),
            ("Skills Missing", ", ".join(missing_skills) if missing_skills else "None — full match"),
            ("Key Projects",   proj_str),
            # ── Quality score fields ──────────────────────────────────────────
            ("Resume Quality",   f"{qs}/100 — {qg}" if qs != "N/A" else "N/A"),
            ("Quality: Passed",  qpass),
            ("Quality: Missing", qfail),
        ]

        info_table = doc.add_table(rows=len(fields), cols=2)
        info_table.style = "Table Grid"
        info_table.columns[0].width = Cm(3.8)
        info_table.columns[1].width = Cm(10.7)

        for row_idx, (label, value) in enumerate(fields):
            lc = info_table.rows[row_idx].cells[0]
            vc = info_table.rows[row_idx].cells[1]
            bg = "FFFFFF" if row_idx % 2 == 0 else "F5F7FA"

            # Highlight quality rows slightly
            if label.startswith("Quality"):
                bg = "FFF8E1" if row_idx % 2 == 0 else "FFF3CD"

            lc.text = label
            _set_cell_bg(lc, "EBF2FA")
            _set_cell_border(lc, "CCCCCC")
            lp = lc.paragraphs[0]
            lp.runs[0].bold = True
            lp.runs[0].font.size = Pt(9)
            lp.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

            vc.text = value
            _set_cell_bg(vc, bg)
            _set_cell_border(vc, "CCCCCC")
            vp = vc.paragraphs[0]
            vp.runs[0].font.size = Pt(9)
            # Colour quality score cell by grade
            if label == "Resume Quality" and qg:
                hx = _quality_color(qg)
                vp.runs[0].font.color.rgb = RGBColor(
                    int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
                )
                vp.runs[0].bold = True
            else:
                vp.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        # Quality improvement tips box
        tips = c.get("quality_tips", [])
        if tips:
            tips_para = doc.add_paragraph()
            tips_para.paragraph_format.space_before = Pt(6)
            tips_para.paragraph_format.space_after  = Pt(4)
            tips_run = tips_para.add_run("Resume Improvement Tips (for candidate feedback):")
            tips_run.bold = True
            tips_run.font.size = Pt(8.5)
            tips_run.font.color.rgb = RGBColor(0x88, 0x55, 0x00)
            for tip in tips:
                tip_p = doc.add_paragraph(style="List Bullet")
                tip_p.paragraph_format.space_after = Pt(2)
                tip_r = tip_p.add_run(tip)
                tip_r.font.size = Pt(8.5)
                tip_r.font.color.rgb = RGBColor(0x55, 0x33, 0x00)

        if i < len(top_candidates) - 1:
            doc.add_paragraph()
            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(10)
            sep_r = sep.add_run("─" * 95)
            sep_r.font.size = Pt(7)
            sep_r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── Section 3: Resume Quality Summary ────────────────────────────────────
    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.paragraph_format.space_after = Pt(10)
    s3_sep = sep2.add_run("─" * 95)
    s3_sep.font.size = Pt(7)
    s3_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    sec3 = doc.add_paragraph()
    sec3.paragraph_format.space_after = Pt(8)
    s3r = sec3.add_run("SECTION 3 — Resume Quality Overview")
    s3r.bold = True
    s3r.font.size = Pt(11)
    s3r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    note3 = doc.add_paragraph()
    note3.paragraph_format.space_after = Pt(10)
    n3r = note3.add_run(
        "Resume Quality Score measures document structure, not job fit. "
        "A candidate with a low quality score but high match score may need resume coaching "
        "before being presented to the hiring panel."
    )
    n3r.italic = True
    n3r.font.size = Pt(9)
    n3r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Quality comparison table
    q_cols   = ["Candidate", "Match Score", "Quality Score", "Grade", "Biggest Gap"]
    q_widths = [4.0, 2.2, 2.2, 2.0, 4.0]
    q_table  = doc.add_table(rows=1, cols=len(q_cols))
    q_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    q_table.style = "Table Grid"

    for i, w in enumerate(q_widths):
        q_table.columns[i].width = int(Cm(w))

    qhrow = q_table.rows[0]
    for i, col in enumerate(q_cols):
        cell = qhrow.cells[i]
        cell.text = col
        _set_cell_bg(cell, "D6E4F0")
        _set_cell_border(cell, "AAAAAA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    for i, c in enumerate(top_candidates):
        qs    = c.get("quality_score", "—")
        qg    = c.get("quality_grade", "—")
        qfail = c.get("quality_failed", [])
        gap   = qfail[0] if qfail else "None — well structured"
        vals  = [
            c.get("name", "?"),
            str(c.get("score", "?")),
            f"{qs}/100" if qs != "—" else "—",
            qg,
            gap,
        ]
        bg  = "FFFFFF" if i % 2 == 0 else "F5F7FA"
        row = q_table.add_row()
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            _set_cell_bg(cell, bg)
            _set_cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(9)
            if j == 3 and qg and qg != "—":
                hx = _quality_color(qg)
                run.font.color.rgb = RGBColor(
                    int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
                )
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
